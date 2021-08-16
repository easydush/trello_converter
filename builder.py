from docx import Document

from models import Card

output_file = 'trello_card.doc'

document = Document()


def head(card: Card):
    document.add_heading('Document Title', 0)


def upload(card: Card):
    print('В файле должно быть написано ', card.name)
    document.add_paragraph(f'{card.name}')
    document.save(output_file)

