import datetime

from docx import Document
import os
from dotenv import load_dotenv

output_file = 'trello'

load_dotenv()
HEADING = os.getenv('HEADING')
SUBHEADING = '...'
PLACE = 'Место проведения:...'
DATE = 'Дата проведения:'
OWNER = os.getenv('OWNER')
PARTICIPANTS = os.getenv('PARTICIPANTS')

TABLE_HEADING = 'Повестка дня:\n'
BOTTOM = os.getenv('BOTTOM')


class DocumentGenerator:
    def __init__(self, card):
        self.card = card
        self.document = Document()
        self.style = self.document.styles['Normal']
        self.font = self.style.font
        self.font.name = 'TimesNewRoman'

    def head(self):
        self.document.add_heading(HEADING, 3)
        self.document.add_heading(SUBHEADING, 5)

    def meta(self):
        self.document.add_paragraph(PLACE)
        self.document.add_paragraph(f'{DATE} {datetime.date.today().__format__("%d.%m.%Y г.")}')
        self.document.add_paragraph(OWNER)
        self.document.add_paragraph(PARTICIPANTS)

    def content(self, checklists):
        self.document.add_paragraph(TABLE_HEADING)
        table = self.document.add_table(rows=1, cols=4)
        table.style = 'TableGrid'
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Организация'
        heading_cells[1].text = 'Замечание'
        heading_cells[2].text = 'Дата'
        heading_cells[3].text = 'Статус'
        for check in checklists:
            for item, value in dict(check).items():
                for point in value:
                    row = table.add_row()
                    row.cells[0].text = item
                    row.cells[1].text = point[0]
                    row.cells[2].text = point[1]
                    row.cells[3].text = point[2]

    def save(self):
        self.document.add_paragraph(BOTTOM)
        self.document.save(f'{self.card.name}.doc')
